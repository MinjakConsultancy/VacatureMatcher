#!/usr/bin/env python3
"""Extract CV docx to plain text for matching."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def extract_docx(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines: list[str] = []
    for p in doc.paragraphs:
        lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: cv_extract.py <cv.docx> [output.txt]", file=sys.stderr)
        return 1

    docx_path = Path(sys.argv[1]).resolve()
    if not docx_path.exists():
        print(f"Not found: {docx_path}", file=sys.stderr)
        return 1

    out_path = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else docx_path.with_suffix(".txt")
    text = extract_docx(docx_path)
    out_path.write_text(text, encoding="utf-8")
    print(f"Wrote {len(text)} chars -> {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
