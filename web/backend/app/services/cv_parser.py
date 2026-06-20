from __future__ import annotations

import io

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError


def _parse_pdf(data: bytes) -> str:
    # Text-based PDFs only; scanned/image PDFs need OCR.
    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as exc:
        raise ValueError("PDF kon niet worden gelezen") from exc

    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)

    result = "\n\n".join(pages).strip()
    if not result:
        raise ValueError("PDF bevat geen leesbare tekst (mogelijk gescande afbeelding)")
    return result


def parse_cv_bytes(data: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace")
    if name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        lines: list[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    lines.append("\t".join(cells))
        return "\n".join(lines)
    if name.endswith(".pdf"):
        return _parse_pdf(data)
    raise ValueError("Alleen .pdf, .docx of .txt toegestaan")
